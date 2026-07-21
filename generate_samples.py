import os
from pathlib import Path

# Create directories
SAMPLE_DATA_DIR = Path("sample_data")
JD_DIR = SAMPLE_DATA_DIR / "job_descriptions"
RESUME_DIR = SAMPLE_DATA_DIR / "resumes"

JD_DIR.mkdir(parents=True, exist_ok=True)
RESUME_DIR.mkdir(parents=True, exist_ok=True)

# 1. Write Job Descriptions
jds = {
    "software_engineer.txt": """Job Title: Senior Software Engineer (Python)
Location: Remote / Hybrid
Experience: 4+ years
Required Skills: Python, FastAPI, PostgreSQL, Docker, Git, REST APIs, System Design.
Preferred Skills: AWS, React, Kubernetes, CI/CD.
Education: Bachelor's degree in Computer Science, Engineering, or related field.
Role Description:
We are looking for a Senior Software Engineer to build and maintain high-performance, scalable backend services. You will design APIs, optimize database queries, containerize services using Docker, and collaborate with frontend teams. You should have strong experience with Python and FastAPI, database design, and cloud deployments.""",

    "product_manager.txt": """Job Title: Product Manager
Location: New York, NY / Hybrid
Experience: 3+ years
Required Skills: Product Strategy, Agile/Scrum, SQL, User Research, Product Roadmap, Market Analysis.
Preferred Skills: Jira, A/B Testing, Data Analytics (Tableau/Mixpanel), Technical Background.
Education: Bachelor's degree in Business, Computer Science, or related; MBA is a plus.
Role Description:
We are looking for a Product Manager to lead the development of our customer-facing platform. You will work closely with engineering, design, and marketing to define product requirements, build roadmaps, analyze user data, and run Scrum meetings. A data-driven approach and excellent communication skills are required.""",

    "data_scientist.txt": """Job Title: Data Scientist
Location: San Francisco, CA / Hybrid
Experience: 3+ years
Required Skills: Python, Machine Learning, SQL, PyTorch/TensorFlow, Pandas/NumPy, Data Visualization.
Preferred Skills: Big Data (Spark), Cloud (GCP/AWS), LLMs/Prompt Engineering, PhD.
Education: Master's or PhD in Statistics, Computer Science, Mathematics, or a quantitative field.
Role Description:
We are seeking a Data Scientist to build predictive models, run statistical analyses, and integrate machine learning features into our product suite. You will analyze large datasets, design and train deep learning models, and work with product teams to translate model outputs into business value."""
}

for name, content in jds.items():
    with open(JD_DIR / name, "w", encoding="utf-8") as f:
        f.write(content.strip())

# 2. Setup Resume Content
resumes_data = [
    # Software Engineers
    {
        "filename": "alex_chen_resume",
        "name": "Alex Chen",
        "email": "alex.chen@email.com",
        "phone": "+1 (555) 019-2834",
        "education": "BS in Computer Science, Stanford University (2020)",
        "experience": "5 years of software engineering experience.\n- Senior Backend Engineer at TechCorp (2022 - Present): Designed and built microservices using Python and FastAPI. Optimized PostgreSQL query times by 40%. Implemented Docker containerization and Kubernetes orchestration.\n- Software Developer at AppStart (2020 - 2022): Developed REST APIs using Flask and worked on frontend integrations using React. Maintained Git workflows.",
        "skills": "Python, FastAPI, PostgreSQL, Docker, Kubernetes, Git, React, REST APIs, AWS, System Design",
        "role": "software_engineer"
    },
    {
        "filename": "emily_watson_resume",
        "name": "Emily Watson",
        "email": "emily.w@email.com",
        "phone": "+1 (555) 021-3948",
        "education": "Self-taught Software Engineer (Online Bootcamps & Certifications)",
        "experience": "3 years of experience.\n- Python Developer at WebSolutions (2023 - Present): Built web scrapers and backend APIs using Django and FastAPI. Maintained PostgreSQL databases.\n- Junior Backend Developer at DevShop (2021 - 2023): Wrote Python scripts for data processing and assisted in REST API development with Flask.",
        "skills": "Python, Django, FastAPI, PostgreSQL, Git, Flask, REST APIs, Web Scraping",
        "role": "software_engineer"
    },
    {
        "filename": "david_miller_resume",
        "name": "David Miller",
        "email": "david.miller@email.com",
        "phone": "+1 (555) 033-9182",
        "education": "BS in Electrical Engineering, Purdue University (2018)",
        "experience": "8 years of software development experience.\n- Senior Systems Developer at HeavyIndustry Corp (2018 - Present): Developed high-performance embedded systems and desktop tools using C++ and Java. Designed relational databases. Managed Git branches.\n- Software Engineer Intern at NetCorp (2017): Developed Java desktop utilities.",
        "skills": "C++, Java, SQL, Git, Linux, Embedded Systems, Multi-threading",
        "role": "software_engineer"
    },
    {
        "filename": "john_doe_resume",
        "name": "John Doe",
        "email": "johndoe@email.com",
        "phone": "+1 (555) 044-8392",
        "education": "High School Diploma",
        "experience": "6 months of frontend experience.\n- Intern Frontend Developer at Pixels Co (2025 - Present): Created landing pages using HTML, CSS, and basic JavaScript. Fixed UI layout alignment issues.",
        "skills": "HTML, CSS, JavaScript, Web Design",
        "role": "software_engineer"
    },

    # Product Managers
    {
        "filename": "sarah_jenkins_resume",
        "name": "Sarah Jenkins",
        "email": "sarah.j@email.com",
        "phone": "+1 (555) 012-7483",
        "education": "MBA from NYU Stern School of Business (2020); BS in Business Administration, Boston University (2016)",
        "experience": "6 years of Product Management experience.\n- Product Manager at Fintech Global (2022 - Present): Formulated product strategy and roadmap for mobile banking app. Led Scrum meetings with 10+ engineers. Wrote SQL queries to analyze user conversion funnels.\n- Associate Product Manager at ShopOnline (2020 - 2022): Conducted user research and managed backlog priorities in Jira. Ran A/B tests that boosted user signups by 15%.",
        "skills": "Product Strategy, Product Roadmap, Agile/Scrum, SQL, User Research, Jira, A/B Testing, Business Analysis",
        "role": "product_manager"
    },
    {
        "filename": "michael_chang_resume",
        "name": "Michael Chang",
        "email": "m.chang@email.com",
        "phone": "+1 (555) 024-8193",
        "education": "BS in Computer Science, UC Berkeley (2021)",
        "experience": "3 years of product/QA experience.\n- Technical Product Manager at DataFlow (2023 - Present): Managed product backlog and roadmap for API integrations. Worked with engineering teams in Scrum sprints. Utilized Jira and SQL for ticket prioritizing and data checks.\n- QA Lead at DataFlow (2021 - 2023): Designed test suites and verified API behavior.",
        "skills": "Product Roadmap, Agile/Scrum, SQL, Jira, QA Testing, API Integration, Technical Communication",
        "role": "product_manager"
    },
    {
        "filename": "sophia_martinez_resume",
        "name": "Sophia Martinez",
        "email": "sophia.m@email.com",
        "phone": "+1 (555) 035-7281",
        "education": "BA in Marketing, Texas A&M University (2022)",
        "experience": "2 years of marketing & product experience.\n- Marketing Specialist & Assistant PM at RetailBoost (2023 - Present): Ran email marketing campaigns and assisted the lead PM in user research and feedback collection. Defined feature specifications for checkout flow improvements.\n- Marketing Intern at AdAgency (2022): Analyzed social media campaigns.",
        "skills": "Marketing Strategy, User Research, Product Support, Copywriting, Social Media, Tableau",
        "role": "product_manager"
    },
    {
        "filename": "robert_vance_resume",
        "name": "Robert Vance",
        "email": "bob.vance@refrigeration.com",
        "phone": "+1 (555) 046-2910",
        "education": "BA in Communications, Scranton University (2015)",
        "experience": "10 years in sales.\n- Sales Director at Vance Refrigeration (2015 - Present): Managed corporate accounts, negotiated refrigeration sales, and achieved 120% of sales quota annually. Led a team of 4 sales representatives.",
        "skills": "B2B Sales, Negotiation, Account Management, Public Speaking, Cold Calling",
        "role": "product_manager"
    },

    # Data Scientists
    {
        "filename": "aris_thorne_resume",
        "name": "Dr. Aris Thorne",
        "email": "aris.thorne@email.com",
        "phone": "+1 (555) 013-8492",
        "education": "PhD in Mathematical Statistics, Columbia University (2022); MS in Applied Mathematics, MIT (2019)",
        "experience": "4 years of Data Science experience.\n- Senior Data Scientist at BioML Labs (2022 - Present): Developed predictive machine learning models using Python and PyTorch. Built data visualization dashboards. Scaled data pipelines on GCP.\n- Research Assistant at Columbia University (2019 - 2022): Researched deep neural architectures and published in top conferences (NeurIPS/ICML).",
        "skills": "Python, Machine Learning, Deep Learning, PyTorch, SQL, Pandas, NumPy, GCP, Data Visualization, R",
        "role": "data_scientist"
    },
    {
        "filename": "elena_rostova_resume",
        "name": "Elena Rostova",
        "email": "elena.r@email.com",
        "phone": "+1 (555) 025-9201",
        "education": "MS in Computer Science, Georgia Tech (2023)",
        "experience": "2 years of experience.\n- Data Scientist at LogisticsCorp (2023 - Present): Built predictive algorithms for shipping delays using Python, Pandas, and Scikit-Learn. Optimized SQL queries on warehouse databases.\n- Graduate Research Assistant at Georgia Tech (2022 - 2023): Analyzed spatial datasets and built regression models.",
        "skills": "Python, Machine Learning, SQL, Pandas, Scikit-Learn, Tableau, Data Analytics",
        "role": "data_scientist"
    },
    {
        "filename": "james_wilson_resume",
        "name": "James Wilson",
        "email": "james.wilson@email.com",
        "phone": "+1 (555) 036-1029",
        "education": "BS in Accounting, Penn State (2018)",
        "experience": "7 years of accounting experience.\n- Senior Accountant at TaxHelpers (2021 - Present): Prepared tax filings, audit reports, and managed corporate ledgers.\n- Completed 1-week Python Data Science Bootcamp (2025).\n- Accountant at LedgerCo (2018 - 2021): Managed payroll and bookkeeping.",
        "skills": "Accounting, Bookkeeping, Excel, Auditing, Financial Analysis, Basic Python",
        "role": "data_scientist"
    },
    {
        "filename": "jane_smith_resume",
        "name": "Jane Smith",
        "email": "janesmith@email.com",
        "phone": "+1 (555) 020-1928",
        "education": "BS in Information Technology, Arizona State University (2019)",
        "experience": "4 years of software development experience.\n- Python Backend Engineer at CloudTech (2021 - Present): Developed REST APIs with Django and FastAPI. Set up CI/CD pipelines and managed serverless functions in AWS.\n- Software Developer at TechSolutions (2019 - 2021): Wrote Python automation scripts and designed relational databases in PostgreSQL.",
        "skills": "Python, Django, FastAPI, PostgreSQL, AWS, CI/CD, Git, REST APIs",
        "role": "software_engineer"
    }
]

def make_txt(resume, path):
    content = f"""
=========================================
CANDIDATE: {resume['name']}
Email: {resume['email']} | Phone: {resume['phone']}
=========================================

EDUCATION:
{resume['education']}

WORK EXPERIENCE:
{resume['experience']}

TECHNICAL SKILLS:
{resume['skills']}
"""
    with open(path.with_suffix(".txt"), "w", encoding="utf-8") as f:
        f.write(content.strip())
    print(f"Created TXT: {path.with_suffix('.txt')}")

def make_docx(resume, path):
    try:
        import docx
        doc = docx.Document()
        doc.add_heading(resume['name'], 0)
        p = doc.add_paragraph()
        p.add_run(f"Email: {resume['email']} | Phone: {resume['phone']}").italic = True
        
        doc.add_heading('Education', level=1)
        doc.add_paragraph(resume['education'])
        
        doc.add_heading('Work Experience', level=1)
        doc.add_paragraph(resume['experience'])
        
        doc.add_heading('Technical Skills', level=1)
        doc.add_paragraph(resume['skills'])
        
        doc.save(path.with_suffix(".docx"))
        print(f"Created DOCX: {path.with_suffix('.docx')}")
    except ImportError:
        # Fallback to text file if docx is not installed
        print("docx library not installed. Falling back to TXT for docx sample.")
        make_txt(resume, path)

def make_pdf(resume, path):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        
        pdf_path = path.with_suffix(".pdf")
        doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Add custom style for spacing
        normal_style = styles["Normal"]
        body_style = ParagraphStyle(
            'Body',
            parent=normal_style,
            spaceAfter=10
        )
        
        story = []
        
        # Header
        story.append(Paragraph(f"<b><font size=18>{resume['name']}</font></b>", styles["Title"]))
        story.append(Paragraph(f"Email: {resume['email']} | Phone: {resume['phone']}", styles["Normal"]))
        story.append(Spacer(1, 15))
        
        # Education
        story.append(Paragraph("<b>Education</b>", styles["Heading2"]))
        story.append(Paragraph(resume['education'].replace('\n', '<br/>'), body_style))
        story.append(Spacer(1, 10))
        
        # Experience
        story.append(Paragraph("<b>Work Experience</b>", styles["Heading2"]))
        story.append(Paragraph(resume['experience'].replace('\n', '<br/>'), body_style))
        story.append(Spacer(1, 10))
        
        # Skills
        story.append(Paragraph("<b>Technical Skills</b>", styles["Heading2"]))
        story.append(Paragraph(resume['skills'].replace('\n', '<br/>'), body_style))
        
        doc.build(story)
        print(f"Created PDF: {pdf_path}")
    except ImportError:
        # Fallback if reportlab is not installed
        print("reportlab library not installed. Falling back to TXT for pdf sample.")
        make_txt(resume, path)

# Iterate over resumes and make them in alternating formats (4 of each)
for i, resume in enumerate(resumes_data):
    filepath = RESUME_DIR / resume['filename']
    if i % 3 == 0:
        make_txt(resume, filepath)
    elif i % 3 == 1:
        make_docx(resume, filepath)
    else:
        make_pdf(resume, filepath)

print("\nSample generation complete!")
