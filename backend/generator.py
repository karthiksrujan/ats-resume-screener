import os
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor

# Setup temp storage directory
from backend.config import UPLOAD_DIR

def clean_filename(name: str) -> str:
    """Sanitize candidate name to make it a safe filename."""
    return "".join(c for c in name if c.isalnum() or c in (" ", "-", "_")).strip().replace(" ", "_")

def generate_ats_pdf(data: dict) -> Path:
    """
    Generates a highly compliant single-column ATS-friendly PDF.
    Uses standard Helvetica, explicit headers, and simple horizontal rules.
    """
    name = data.get("name", "Candidate")
    safe_name = clean_filename(name)
    pdf_path = UPLOAD_DIR / f"{safe_name}_ATS_Resume.pdf"
    
    # Page setup - 0.75 inch (54 pt) margins for standard layout
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    # Primary color - solid dark slate for text (Hex #1e293b)
    text_color = HexColor("#1e293b")
    
    name_style = ParagraphStyle(
        'ATS_Name',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        alignment=1, # Centered
        textColor=text_color,
        spaceAfter=4
    )
    
    contact_style = ParagraphStyle(
        'ATS_Contact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        alignment=1, # Centered
        textColor=HexColor("#475569"),
        spaceAfter=15
    )
    
    section_heading = ParagraphStyle(
        'ATS_SectionHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=text_color,
        spaceBefore=10,
        spaceAfter=4
    )
    
    body_style = ParagraphStyle(
        'ATS_Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=text_color,
        spaceAfter=6
    )
    
    bold_body = ParagraphStyle(
        'ATS_BoldBody',
        parent=body_style,
        fontName='Helvetica-Bold'
    )
    
    bullet_style = ParagraphStyle(
        'ATS_Bullet',
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=3
    )

    story = []
    
    # 1. Header (Name & Contact info)
    story.append(Paragraph(name.upper(), name_style))
    
    contact_parts = []
    if data.get("email"): contact_parts.append(data["email"])
    if data.get("phone"): contact_parts.append(data["phone"])
    if data.get("location"): contact_parts.append(data["location"])
    if data.get("linkedin"): contact_parts.append(data["linkedin"])
    
    story.append(Paragraph(" | ".join(contact_parts), contact_style))
    
    # Helper to add section headers with divider lines
    def add_section_header(title: str):
        story.append(Spacer(1, 4))
        story.append(Paragraph(title.upper(), section_heading))
        story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#cbd5e1"), spaceBefore=2, spaceAfter=8))

    # 2. Professional Summary
    if data.get("summary"):
        add_section_header("Professional Summary")
        story.append(Paragraph(data["summary"], body_style))
        
    # 3. Work Experience
    experience = data.get("experience", [])
    if experience:
        add_section_header("Work Experience")
        for job in experience:
            # Job Title and Company line
            job_title = job.get("title", "Position")
            company = job.get("company", "Company")
            dates = job.get("dates", "")
            loc = job.get("location", "")
            
            header_text = f"<b>{job_title}</b> — {company}"
            if loc:
                header_text += f", {loc}"
                
            story.append(Paragraph(f"<font size=10>{header_text}</font>", body_style))
            if dates:
                story.append(Paragraph(f"<i>{dates}</i>", ParagraphStyle('ATS_Date', parent=body_style, fontName='Helvetica-Oblique', spaceAfter=4)))
            
            # Bullet points
            description = job.get("description", "")
            if description:
                # Split by bullet indicator or newline
                bullets = [b.strip() for b in description.split("\n") if b.strip()]
                for b in bullets:
                    # Clean leading bullet characters if present
                    if b.startswith("-") or b.startswith("•"):
                        b = b[1:].strip()
                    story.append(Paragraph(f"&bull; {b}", bullet_style))
            story.append(Spacer(1, 5))
            
    # 4. Education
    education = data.get("education", [])
    if education:
        add_section_header("Education")
        for edu in education:
            degree = edu.get("degree", "Degree")
            school = edu.get("school", "Institution")
            dates = edu.get("dates", "")
            loc = edu.get("location", "")
            
            edu_text = f"<b>{degree}</b> — {school}"
            if loc:
                edu_text += f", {loc}"
            story.append(Paragraph(edu_text, body_style))
            
            if dates:
                story.append(Paragraph(f"<i>{dates}</i>", ParagraphStyle('ATS_EduDate', parent=body_style, fontName='Helvetica-Oblique', spaceAfter=4)))
            story.append(Spacer(1, 3))
            
    # 5. Skills
    skills = data.get("skills", "")
    if skills:
        add_section_header("Skills & Competencies")
        story.append(Paragraph(skills, body_style))

    # Build the document
    doc.build(story)
    return pdf_path

def generate_ats_docx(data: dict) -> Path:
    """
    Generates a highly structured DOCX file using python-docx.
    Applies single-column layout, plain styling, and clean spacing.
    """
    name = data.get("name", "Candidate")
    safe_name = clean_filename(name)
    docx_path = UPLOAD_DIR / f"{safe_name}_ATS_Resume.docx"
    
    doc = Document()
    
    # Page setup - 0.75" margins
    sections = doc.sections
    for section in sections:
        section.top_margin = inch * 0.75
        section.bottom_margin = inch * 0.75
        section.left_margin = inch * 0.75
        section.right_margin = inch * 0.75
        
    # Header Info
    p_name = doc.add_paragraph()
    p_name.alignment = 1 # Center
    run_name = p_name.add_run(name.upper())
    run_name.font.name = 'Arial'
    run_name.font.size = docx.shared.Pt(18)
    run_name.bold = True
    
    contact_parts = []
    if data.get("email"): contact_parts.append(data["email"])
    if data.get("phone"): contact_parts.append(data["phone"])
    if data.get("location"): contact_parts.append(data["location"])
    if data.get("linkedin"): contact_parts.append(data["linkedin"])
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = 1 # Center
    run_contact = p_contact.add_run(" | ".join(contact_parts))
    run_contact.font.name = 'Arial'
    run_contact.font.size = docx.shared.Pt(9.5)
    
    # Helper to add sections
    def add_section_title(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = docx.shared.Pt(11)
        p.paragraph_format.space_before = docx.shared.Pt(12)
        p.paragraph_format.space_after = docx.shared.Pt(4)
        p.paragraph_format.keep_with_next = True
        
        # Add thin bottom border using a simple text separator
        p_line = doc.add_paragraph()
        run_line = p_line.add_run("―" * 60)
        run_line.font.color.rgb = docx.shared.RGBColor(203, 213, 225)
        run_line.font.size = docx.shared.Pt(6)
        p_line.paragraph_format.space_after = docx.shared.Pt(6)

    # 1. Summary
    if data.get("summary"):
        add_section_title("Professional Summary")
        p = doc.add_paragraph()
        run = p.add_run(data["summary"])
        run.font.name = 'Arial'
        run.font.size = docx.shared.Pt(10)
        p.paragraph_format.space_after = docx.shared.Pt(8)
        
    # 2. Experience
    experience = data.get("experience", [])
    if experience:
        add_section_title("Work Experience")
        for job in experience:
            p_job = doc.add_paragraph()
            p_job.paragraph_format.space_before = docx.shared.Pt(4)
            p_job.paragraph_format.space_after = docx.shared.Pt(2)
            
            run_title = p_job.add_run(job.get("title", ""))
            run_title.bold = True
            run_title.font.name = 'Arial'
            run_title.font.size = docx.shared.Pt(10)
            
            p_job.add_run(f" | {job.get('company', '')}")
            if job.get("location"):
                p_job.add_run(f", {job.get('location', '')}")
                
            if job.get("dates"):
                p_date = doc.add_paragraph()
                p_date.paragraph_format.space_after = docx.shared.Pt(4)
                run_date = p_date.add_run(job.get("dates", ""))
                run_date.italic = True
                run_date.font.name = 'Arial'
                run_date.font.size = docx.shared.Pt(9)
                
            desc = job.get("description", "")
            if desc:
                bullets = [b.strip() for b in desc.split("\n") if b.strip()]
                for b in bullets:
                    if b.startswith("-") or b.startswith("•"):
                        b = b[1:].strip()
                    p_b = doc.add_paragraph(style='List Bullet')
                    p_b.paragraph_format.space_after = docx.shared.Pt(2)
                    run_b = p_b.add_run(b)
                    run_b.font.name = 'Arial'
                    run_b.font.size = docx.shared.Pt(9.5)
                    
    # 3. Education
    education = data.get("education", [])
    if education:
        add_section_title("Education")
        for edu in education:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_after = docx.shared.Pt(2)
            run_deg = p_edu.add_run(edu.get("degree", ""))
            run_deg.bold = True
            run_deg.font.name = 'Arial'
            run_deg.font.size = docx.shared.Pt(10)
            
            p_edu.add_run(f" | {edu.get('school', '')}")
            if edu.get("location"):
                p_edu.add_run(f", {edu.get('location', '')}")
                
            if edu.get("dates"):
                p_edate = doc.add_paragraph()
                p_edate.paragraph_format.space_after = docx.shared.Pt(6)
                run_edate = p_edate.add_run(edu.get("dates", ""))
                run_edate.italic = True
                run_edate.font.name = 'Arial'
                run_edate.font.size = docx.shared.Pt(9)

    # 4. Skills
    skills = data.get("skills", "")
    if skills:
        add_section_title("Skills & Competencies")
        p_skills = doc.add_paragraph()
        run_skills = p_skills.add_run(skills)
        run_skills.font.name = 'Arial'
        run_skills.font.size = docx.shared.Pt(10)
        
    doc.save(docx_path)
    return docx_path

def generate_cover_letter_pdf(data: dict) -> Path:
    """
    Generates a beautifully typeset professional PDF cover letter.
    """
    name = data.get("name", "Candidate")
    email = data.get("email", "N/A")
    phone = data.get("phone", "N/A")
    letter_text = data.get("text", "")
    
    safe_name = clean_filename(name)
    pdf_path = UPLOAD_DIR / f"{safe_name}_Cover_Letter.pdf"
    
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    text_color = HexColor("#1e293b")
    
    name_style = ParagraphStyle(
        'CL_Name',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=text_color,
        spaceAfter=4
    )
    
    contact_style = ParagraphStyle(
        'CL_Contact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=HexColor("#475569"),
        spaceAfter=15
    )
    
    body_style = ParagraphStyle(
        'CL_Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=text_color,
        spaceAfter=12
    )
    
    story = []
    
    # Header block
    story.append(Paragraph(name, name_style))
    contact_info = f"Email: {email} | Phone: {phone}"
    story.append(Paragraph(contact_info, contact_style))
    story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#cbd5e1"), spaceBefore=0, spaceAfter=20))
    
    # Body text
    # Split text into paragraphs by double-newline
    paragraphs = [p.strip() for p in letter_text.split('\n\n') if p.strip()]
    for p_text in paragraphs:
        # Format single newlines into spaces, keeping spacing clean
        clean_p_text = " ".join(p_text.split('\n'))
        story.append(Paragraph(clean_p_text, body_style))
        
    doc.build(story)
    return pdf_path
