import os
from pathlib import Path
from pypdf import PdfReader
import docx

def parse_txt(file_path: Path) -> str:
    """Extract text from plain text files."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception as e:
        raise ValueError(f"Error reading TXT file {file_path.name}: {str(e)}")

def parse_pdf(file_path: Path) -> str:
    """Extract text from PDF files using pypdf."""
    try:
        reader = PdfReader(file_path)
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n\n".join(text_parts).strip()
    except Exception as e:
        raise ValueError(f"Error reading PDF file {file_path.name}: {str(e)}")

def parse_docx(file_path: Path) -> str:
    """Extract text from DOCX files using python-docx."""
    try:
        doc = docx.Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
                
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
        return "\n\n".join(text_parts).strip()
    except Exception as e:
        raise ValueError(f"Error reading DOCX file {file_path.name}: {str(e)}")

def parse_resume(file_path: Path | str) -> str:
    """
    Main parser entrypoint. Resolves the file format from extension
    and extracts text.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
        
    ext = path.suffix.lower()
    
    if ext == ".txt":
        return parse_txt(path)
    elif ext == ".pdf":
        return parse_pdf(path)
    elif ext == ".docx":
        return parse_docx(path)
    else:
        # Fallback to plain text read if unknown format
        try:
            return parse_txt(path)
        except Exception:
            raise ValueError(f"Unsupported file format: {ext}")
